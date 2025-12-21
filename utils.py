# /ai_detector_pro/utils.py
# VERSION: 2.1 – Phase 2.1 + 2.2 Stable Utilities

import io
import re
import json
import hashlib
from datetime import datetime
from io import BytesIO

import html
import difflib

# --------------------------------------------------
# OPTIONAL DEPENDENCIES
# --------------------------------------------------
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    import docx
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import parse_xml
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ==================================================
# SECTION 1: FILE READING
# ==================================================
def read_file_content(uploaded_file):
    """Reads content from txt, pdf, or docx uploads."""
    try:
        uploaded_file.seek(0)
        ext = uploaded_file.name.split(".")[-1].lower()
        data = uploaded_file.read()

        if ext in ["txt", "md"]:
            return data.decode("utf-8")

        if ext == "docx":
            if not DOCX_AVAILABLE:
                return "ERROR: python-docx not installed."
            document = docx.Document(io.BytesIO(data))
            return "\n".join(p.text for p in document.paragraphs)

        if ext == "pdf":
            if not fitz:
                return "ERROR: PyMuPDF not installed."
            with fitz.open(stream=data, filetype="pdf") as pdf:
                text = "".join(page.get_text("text") for page in pdf)
                return text.strip() or "ERROR: PDF contains no extractable text."

        return "ERROR: Unsupported file type."

    except Exception as e:
        return f"ERROR: Could not read file ({e})"

# ==================================================
# SECTION 2: EXPORTS
# ==================================================
def create_docx_export(results):
    """Creates DOCX report with AI highlights."""
    if not DOCX_AVAILABLE:
        return None

    document = docx.Document()
    document.add_heading("AI Writing Analysis Report", 0)

    document.add_paragraph(
        f"Human Score: {results['composite_human_score']}%"
    )
    document.add_paragraph(
        f"Generated: {datetime.now().strftime('%B %d, %Y')}"
    )

    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Flag"
    hdr[1].text = "Sentence"

    color_map = {
        "AI_VERY_PREDICTABLE": "FFD6D6",
        "AI_PREDICTABLE": "FFF0CC",
        "BOILERPLATE": "E0E0FF",
        "HUMAN": "FFFFFF",
    }

    for item in results["sentence_analysis"]:
        row = table.add_row().cells
        sd = item["suggestion_data"]
        row[0].text = f"{sd['symbol']} {sd['short']}"
        row[1].text = item["sentence"]

        color = color_map.get(item["flag"], "FFFFFF")
        shade = parse_xml(
            f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{color}"/>'
        )
        row[1]._tc.get_or_add_tcPr().append(shade)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.read()

def create_html_export(results):
    """Creates HTML analysis report."""
    highlights = []
    for item in results["sentence_analysis"]:
        color = {
            "AI_VERY_PREDICTABLE": "#ffdcdc",
            "AI_PREDICTABLE": "#fff0cc",
            "BOILERPLATE": "#f0f0ff",
            "HUMAN": "transparent",
        }.get(item["flag"], "transparent")

        highlights.append(
            f'<span style="background:{color};padding:3px;border-radius:4px;">{html.escape(item["sentence"])}</span>'
        )

    return f"""
    <html>
    <body style="font-family:sans-serif;padding:20px;">
        <h1>AI Writing Analysis</h1>
        <p><b>Human Score:</b> {results['composite_human_score']}%</p>
        <div style="line-height:1.8;">{" ".join(highlights)}</div>
    </body>
    </html>
    """.encode("utf-8")

def create_json_export(results):
    """JSON export of analysis results."""
    return json.dumps(results, indent=2).encode("utf-8")

# ==================================================
# SECTION 3: UI HELPERS
# ==================================================
def generate_highlighted_text_html(sentence_data):
    """HTML sentence highlights with tooltips."""
    blocks = []
    for item in sentence_data:
        sd = item["suggestion_data"]
        tooltip = f"{sd['short']} (Perp: {item['perplexity']:.1f})"
        color = {
            "AI_VERY_PREDICTABLE": "#ffdcdc",
            "AI_PREDICTABLE": "#fff0cc",
            "BOILERPLATE": "#f0f0ff",
            "HUMAN": "transparent",
        }.get(item["flag"], "transparent")

        blocks.append(
            f'<span title="{tooltip}" style="background:{color};padding:3px;border-radius:4px;">{html.escape(item["sentence"])}</span>'
        )

    return f'<div style="line-height:1.8;">{" ".join(blocks)}</div>'

def highlight_sentence_diff(original_text, improved_text):
    """Sentence-level diff (proofreading mode)."""
    orig = re.split(r"(?<=[.!?])\s+", original_text.strip())
    new = re.split(r"(?<=[.!?])\s+", improved_text.strip())

    orig_set, new_set = set(orig), set(new)

    def block(s, color=None):
        style = f"background:{color};padding:6px;border-radius:6px;" if color else ""
        return f"<p style='{style}'>{html.escape(s)}</p>"

    orig_html = "".join(
        block(s, "#ffe6e6") if s not in new_set else block(s) for s in orig
    )
    new_html = "".join(
        block(s, "#e6ffe6") if s not in orig_set else block(s) for s in new
    )

    return orig_html, new_html

def highlight_grammar_diff(original: str, corrected: str):
    """Word-level grammar diff."""
    orig_words = original.split()
    new_words = corrected.split()

    diff = difflib.ndiff(orig_words, new_words)

    orig_out, new_out = [], []

    for d in diff:
        tag, word = d[:2], html.escape(d[2:])
        if tag == "- ":
            orig_out.append(f"<span style='background:#ffd6d6'>{word}</span>")
        elif tag == "+ ":
            new_out.append(f"<span style='background:#d6ffd6'>{word}</span>")
        elif tag == "  ":
            orig_out.append(word)
            new_out.append(word)

    return (
        "<div>" + " ".join(orig_out) + "</div>",
        "<div>" + " ".join(new_out) + "</div>",
    )

def generate_rewrite_suggestions(sentence, flag):
    """Rule-based rewrite suggestions."""
    if flag in ["AI_VERY_PREDICTABLE", "AI_PREDICTABLE"]:
        return [
            {"strategy": "Add Detail", "rewrite": f"[Example] {sentence}"},
            {"strategy": "Personalize", "rewrite": f"In my experience, {sentence.lower()}"},
        ]
    if flag == "BOILERPLATE":
        cleaned = re.sub(
            r"^(In conclusion|Furthermore|Moreover|Thus|Hence)[,\s]+",
            "",
            sentence,
            flags=re.I,
        )
        return [{"strategy": "Direct Start", "rewrite": cleaned.capitalize()}]

    return [{"strategy": "Human Touch", "rewrite": f"{sentence} — something I’ve seen myself."}]

def get_unique_key(element, idx):
    """Stable Streamlit key generator (reserved for future use)."""
    return hashlib.md5(f"{element[:20]}-{idx}".encode()).hexdigest()
